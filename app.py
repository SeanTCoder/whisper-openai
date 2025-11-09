import streamlit as st
import whisper
import tempfile
import os
import sys
import io
import json
from pathlib import Path
from contextlib import redirect_stdout, redirect_stderr
import threading
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from concurrent.futures import ThreadPoolExecutor, as_completed
import subprocess

# Check and install watchdog if needed (required for Streamlit file watching)
try:
    import watchdog
except ImportError:
    import subprocess
    import sys
    try:
        # Try to install watchdog silently
        subprocess.check_call([sys.executable, "-m", "pip", "install", "watchdog>=3.0.0"], 
                            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        import watchdog
    except Exception as e:
        # If installation fails, continue anyway (watchdog is optional for basic functionality)
        pass

# Page configuration
st.set_page_config(
    page_title="Whisper Speech Transcription",
    page_icon="🎤",
    layout="wide"
)

# Storage directory for transcriptions
STORAGE_DIR = Path("transcriptions")
STORAGE_DIR.mkdir(exist_ok=True)
TRANSCRIPTIONS_FILE = STORAGE_DIR / "transcriptions.json"
ARCHIVED_TRANSCRIPTIONS_FILE = STORAGE_DIR / "archived_transcriptions.json"

# Helper functions for storage
def load_transcriptions():
    """Load stored transcriptions from JSON file"""
    if TRANSCRIPTIONS_FILE.exists():
        try:
            with open(TRANSCRIPTIONS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def save_transcription(transcription_data):
    """Save a transcription to storage"""
    transcriptions = load_transcriptions()
    transcriptions.append(transcription_data)
    with open(TRANSCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(transcriptions, f, indent=2, ensure_ascii=False)

def load_archived_transcriptions():
    """Load archived transcriptions from JSON file"""
    if ARCHIVED_TRANSCRIPTIONS_FILE.exists():
        try:
            with open(ARCHIVED_TRANSCRIPTIONS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def archive_transcription(transcription_id):
    """Move transcription to archive"""
    transcriptions = load_transcriptions()
    archived = load_archived_transcriptions()
    
    # Find and move transcription
    for i, trans in enumerate(transcriptions):
        if trans.get('id') == transcription_id:
            trans['archived_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            archived.append(trans)
            transcriptions.pop(i)
            break
    
    # Save both files
    with open(TRANSCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(transcriptions, f, indent=2, ensure_ascii=False)
    with open(ARCHIVED_TRANSCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(archived, f, indent=2, ensure_ascii=False)
    return True

def purge_transcription(transcription_id, from_archive=False):
    """Permanently delete transcription"""
    if from_archive:
        archived = load_archived_transcriptions()
        archived = [t for t in archived if t.get('id') != transcription_id]
        with open(ARCHIVED_TRANSCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(archived, f, indent=2, ensure_ascii=False)
    else:
        transcriptions = load_transcriptions()
        transcriptions = [t for t in transcriptions if t.get('id') != transcription_id]
        with open(TRANSCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(transcriptions, f, indent=2, ensure_ascii=False)
    return True

def format_timestamp(seconds):
    """Format seconds to HH:MM:SS.mmm or MM:SS.mmm"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    return f"{minutes:02d}:{secs:02d}.{millis:03d}"

# Configuration for large file processing
CHUNK_SIZE_MB = 200  # Maximum chunk size in MB
MAX_WORKERS = 5  # Number of parallel threads
CHUNK_DURATION_SECONDS = 300  # 5 minutes per chunk (reduced to create more chunks for better parallelization)
MIN_CHUNK_DURATION_SECONDS = 60  # Minimum chunk duration (1 minute) to avoid too many tiny chunks

def get_file_duration(file_path):
    """Get the duration of an audio/video file in seconds using ffprobe"""
    try:
        cmd = [
            "ffprobe",
            "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            file_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        duration = float(result.stdout.strip())
        return duration
    except Exception as e:
        # Try alternative method if ffprobe fails
        try:
            # Use ffmpeg to get duration
            cmd = [
                "ffmpeg",
                "-i", file_path,
                "-f", "null",
                "-"
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, stderr=subprocess.STDOUT)
            # Parse duration from ffmpeg output (format: Duration: HH:MM:SS.mmm)
            for line in result.stdout.split('\n'):
                if 'Duration:' in line:
                    time_str = line.split('Duration:')[1].split(',')[0].strip()
                    parts = time_str.split(':')
                    hours = float(parts[0])
                    minutes = float(parts[1])
                    seconds = float(parts[2])
                    return hours * 3600 + minutes * 60 + seconds
        except:
            pass
        return None

def validate_chunk_file(chunk_path, min_duration=1.0):
    """Validate that a chunk file has audio content"""
    try:
        # Check file exists and has content
        if not os.path.exists(chunk_path):
            return False
        
        file_size = os.path.getsize(chunk_path)
        # Minimum file size: at least 1KB (roughly 0.1 seconds of audio)
        if file_size < 1024:
            return False
        
        # Check actual duration using ffprobe
        cmd = [
            "ffprobe",
            "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            chunk_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        chunk_duration = float(result.stdout.strip())
        
        # Chunk must have at least min_duration seconds of audio
        return chunk_duration >= min_duration
    except:
        return False

def split_file_into_chunks(file_path, chunk_duration=CHUNK_DURATION_SECONDS, max_workers=MAX_WORKERS):
    """Split a large file into time-based chunks using ffmpeg
    
    Creates enough chunks to utilize the thread pool effectively.
    Adjusts chunk size to ensure we have at least max_workers chunks
    (or close to it) for better parallelization.
    """
    duration = get_file_duration(file_path)
    if duration is None:
        # Fallback: estimate based on file size (rough estimate)
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        # Rough estimate: 1MB per minute for audio
        duration = file_size_mb * 60
    
    chunks = []
    chunk_dir = tempfile.mkdtemp()
    
    # Calculate optimal chunk duration to utilize thread pool
    # We want at least max_workers chunks, but not too many tiny chunks
    initial_num_chunks = int(duration / chunk_duration) + 1
    
    # If we have fewer chunks than workers, reduce chunk duration to create more chunks
    if initial_num_chunks < max_workers:
        # Adjust chunk duration to create approximately max_workers chunks
        # But don't go below MIN_CHUNK_DURATION_SECONDS
        optimal_chunk_duration = max(
            MIN_CHUNK_DURATION_SECONDS,
            duration / max_workers
        )
        chunk_duration = optimal_chunk_duration
        num_chunks = int(duration / chunk_duration) + 1
    else:
        num_chunks = initial_num_chunks
    
    # Ensure we don't create chunks that are too short
    if duration < MIN_CHUNK_DURATION_SECONDS:
        # File is very short, process as single chunk
        num_chunks = 1
        chunk_duration = duration
    
    for i in range(num_chunks):
        start_time = i * chunk_duration
        end_time = min((i + 1) * chunk_duration, duration)
        chunk_length = end_time - start_time
        
        # Skip chunks that are too short (less than 2 seconds)
        # This prevents empty or near-empty chunks that cause tensor errors
        if chunk_length < 2.0:
            # If this is the last chunk and it's short, merge it with the previous one
            if i == num_chunks - 1 and chunks:
                # Extend the last chunk to include this short segment
                last_chunk = chunks[-1]
                last_chunk["end_time"] = duration
                last_chunk["duration"] = duration - last_chunk["start_time"]
            continue
        
        chunk_path = os.path.join(chunk_dir, f"chunk_{i:04d}.wav")
        
        # Extract chunk using ffmpeg with better error handling
        cmd = [
            "ffmpeg",
            "-i", file_path,
            "-ss", str(start_time),
            "-t", str(chunk_length),
            "-acodec", "pcm_s16le",
            "-ar", "16000",
            "-ac", "1",
            "-avoid_negative_ts", "make_zero",  # Handle timestamp issues
            "-y",  # Overwrite output file
            chunk_path
        ]
        
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            
            # Validate the chunk file
            if validate_chunk_file(chunk_path, min_duration=1.0):
                chunks.append({
                    "path": chunk_path,
                    "index": i,
                    "start_time": start_time,
                    "end_time": end_time,
                    "duration": chunk_length
                })
            else:
                # Chunk is empty or too short, skip it
                try:
                    os.remove(chunk_path)
                except:
                    pass
                continue
                
        except subprocess.CalledProcessError as e:
            # Log the error but continue
            error_msg = e.stderr.decode() if e.stderr else str(e)
            if "Output file does not contain any stream" not in error_msg:
                # Only warn if it's not an expected empty chunk
                pass
            try:
                if os.path.exists(chunk_path):
                    os.remove(chunk_path)
            except:
                pass
            continue
    
    return chunks, chunk_dir

def transcribe_chunk(chunk_info, model, transcribe_options, progress_callback=None):
    """Transcribe a single chunk with validation"""
    try:
        # Validate chunk file before processing
        if not os.path.exists(chunk_info["path"]):
            return {
                "chunk_index": chunk_info["index"],
                "start_time": chunk_info["start_time"],
                "error": "Chunk file does not exist",
                "success": False
            }
        
        file_size = os.path.getsize(chunk_info["path"])
        if file_size < 1024:  # Less than 1KB
            return {
                "chunk_index": chunk_info["index"],
                "start_time": chunk_info["start_time"],
                "error": "Chunk file is too small (likely empty)",
                "success": False
            }
        
        # Try to load audio first to validate it's not empty
        try:
            import whisper
            audio = whisper.load_audio(chunk_info["path"])
            if len(audio) == 0 or audio.size == 0:
                return {
                    "chunk_index": chunk_info["index"],
                    "start_time": chunk_info["start_time"],
                    "error": "Chunk contains no audio data",
                    "success": False
                }
        except Exception as e:
            return {
                "chunk_index": chunk_info["index"],
                "start_time": chunk_info["start_time"],
                "error": f"Failed to load audio: {str(e)}",
                "success": False
            }
        
        # Transcribe the chunk
        result = model.transcribe(chunk_info["path"], **transcribe_options)
        
        # Adjust timestamps to account for chunk offset
        if result.get("segments"):
            for segment in result["segments"]:
                segment["start"] += chunk_info["start_time"]
                segment["end"] += chunk_info["start_time"]
                if segment.get("words"):
                    for word in segment["words"]:
                        word["start"] += chunk_info["start_time"]
                        word["end"] += chunk_info["start_time"]
        
        return {
            "chunk_index": chunk_info["index"],
            "start_time": chunk_info["start_time"],
            "result": result,
            "success": True
        }
    except Exception as e:
        error_msg = str(e)
        # Provide more helpful error messages
        if "reshape" in error_msg and "0 elements" in error_msg:
            error_msg = "Chunk contains no audio data or is empty"
        return {
            "chunk_index": chunk_info["index"],
            "start_time": chunk_info["start_time"],
            "error": error_msg,
            "success": False
        }

def combine_chunk_results(chunk_results, chunks_info):
    """Combine transcription results from multiple chunks, maintaining temporal order"""
    # Sort by chunk index to maintain order
    chunk_results.sort(key=lambda x: x["chunk_index"])
    
    all_segments = []
    all_text_parts = []
    language = None
    total_duration = 0
    processed_chunk_indices = set()
    
    # Create a map of chunk indices to their time ranges
    chunk_time_map = {chunk["index"]: (chunk["start_time"], chunk["end_time"]) for chunk in chunks_info}
    
    for chunk_result in chunk_results:
        if not chunk_result["success"]:
            # Track which chunks were skipped
            chunk_idx = chunk_result["chunk_index"]
            processed_chunk_indices.add(chunk_idx)
            continue
        
        chunk_idx = chunk_result["chunk_index"]
        processed_chunk_indices.add(chunk_idx)
        
        result = chunk_result["result"]
        
        if language is None:
            language = result.get("language")
        
        # Collect segments with proper timestamp adjustment
        if result.get("segments"):
            for segment in result["segments"]:
                # Ensure timestamps are within chunk bounds
                chunk_start = chunk_result["start_time"]
                segment_start = segment.get("start", 0)
                segment_end = segment.get("end", 0)
                
                # Validate timestamps are reasonable
                if segment_start >= 0 and segment_end > segment_start:
                    all_segments.append(segment)
        
        # Collect text - preserve spacing between chunks
        text = result.get("text", "").strip()
        if text:
            all_text_parts.append(text)
        
        # Update total duration based on chunk end time
        chunk_start = chunk_result["start_time"]
        chunk_duration = result.get("duration", 0)
        chunk_end_time = chunk_start + chunk_duration
        total_duration = max(total_duration, chunk_end_time)
    
    # Check for missing chunks
    expected_indices = set(range(len(chunks_info)))
    missing_indices = expected_indices - processed_chunk_indices
    
    # Sort segments by start time to ensure temporal order
    all_segments.sort(key=lambda x: x.get("start", 0))
    
    # Combine text with proper spacing
    combined_text = " ".join(all_text_parts)
    
    # Build full result with all metadata
    combined_result = {
        "text": combined_text,
        "segments": all_segments,
        "language": language,
        "duration": total_duration
    }
    
    # Add metadata about chunk processing
    if missing_indices:
        combined_result["_metadata"] = {
            "missing_chunks": sorted(list(missing_indices)),
            "processed_chunks": sorted(list(processed_chunk_indices)),
            "total_chunks": len(chunks_info)
        }
    
    return combined_result

def create_docx(transcription_text, filename, metadata=None, segments=None):
    """Create a DOCX file from transcription text with timestamps"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Transcription', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata if provided
    if metadata:
        doc.add_paragraph(f"Source File: {metadata.get('filename', 'Unknown')}")
        doc.add_paragraph(f"Language: {metadata.get('language', 'Unknown')}")
        doc.add_paragraph(f"Duration: {metadata.get('duration', 0):.2f} seconds")
        doc.add_paragraph(f"Date: {metadata.get('date', 'Unknown')}")
        doc.add_paragraph("")  # Empty line
    
    # Add transcription text with timestamps if segments are provided
    if segments:
        doc.add_heading('Transcription with Timestamps', level=1)
        for segment in segments:
            start = segment.get('start', 0)
            end = segment.get('end', 0)
            text = segment.get('text', '').strip()
            
            # Format timestamp
            def format_time(seconds):
                hours = int(seconds // 3600)
                minutes = int((seconds % 3600) // 60)
                secs = int(seconds % 60)
                millis = int((seconds % 1) * 1000)
                if hours > 0:
                    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
                return f"{minutes:02d}:{secs:02d}.{millis:03d}"
            
            timestamp = f"[{format_time(start)} --> {format_time(end)}]"
            para = doc.add_paragraph()
            para.add_run(timestamp).bold = True
            para.add_run(f" {text}")
            doc.add_paragraph("")  # Empty line between segments
        
        doc.add_paragraph("")  # Empty line
        doc.add_heading('Plain Text (No Timestamps)', level=1)
    
    # Add plain transcription text
    doc.add_paragraph(transcription_text)
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# Title and description
st.title("🎤 Whisper Speech Transcription")
st.markdown("Upload an audio file (MP3, WAV, FLAC, etc.) to transcribe or translate speech to text.")

# Add tabs for main transcription and history
tab1, tab2 = st.tabs(["🎤 Transcribe", "📚 Stored Transcriptions"])

# Sidebar for settings
with st.sidebar:
    st.header("⚙️ Settings")
    
    # Model selection
    model_options = {
        "tiny": "Tiny (39M params) - Fastest, least accurate",
        "base": "Base (74M params) - Fast, good balance",
        "small": "Small (244M params) - Good accuracy",
        "medium": "Medium (769M params) - Better accuracy",
        "large": "Large (1550M params) - Best accuracy, slowest",
        "turbo": "Turbo (809M params) - Optimized for speed (default)"
    }
    
    selected_model = st.selectbox(
        "Select Model",
        options=list(model_options.keys()),
        index=5,  # Default to turbo
        format_func=lambda x: f"{x} - {model_options[x].split(' - ')[1]}"
    )
    
    st.caption(model_options[selected_model])
    
    # Task selection
    task = st.radio(
        "Task",
        options=["transcribe", "translate"],
        help="Transcribe: Keep original language. Translate: Translate to English."
    )
    
    # Language selection
    language = st.selectbox(
        "Language (optional - leave as 'Auto-detect' to detect automatically)",
        options=["Auto-detect"] + ["English", "Spanish", "French", "German", "Japanese", "Chinese", "Korean", "Italian", "Portuguese", "Russian", "Arabic", "Hindi", "Dutch", "Polish", "Turkish", "Swedish", "Norwegian", "Danish", "Finnish", "Greek", "Czech", "Hungarian", "Romanian", "Thai", "Vietnamese", "Indonesian", "Malay", "Hebrew", "Ukrainian", "Catalan", "Basque", "Galician", "Welsh", "Irish", "Scottish Gaelic", "Breton", "Esperanto", "Latin", "Yiddish", "Afrikaans", "Albanian", "Amharic", "Armenian", "Assamese", "Azerbaijani", "Bashkir", "Belarusian", "Bengali", "Bosnian", "Bulgarian", "Burmese", "Cantonese", "Castilian", "Croatian", "Estonian", "Faroese", "Flemish", "Georgian", "Gujarati", "Haitian", "Haitian Creole", "Hausa", "Hawaiian", "Icelandic", "Javanese", "Kannada", "Kazakh", "Khmer", "Lao", "Latvian", "Letzeburgesch", "Lingala", "Lithuanian", "Luxembourgish", "Macedonian", "Malagasy", "Malayalam", "Maltese", "Mandarin", "Maori", "Marathi", "Moldavian", "Moldovan", "Mongolian", "Myanmar", "Nepali", "Nynorsk", "Occitan", "Panjabi", "Pashto", "Persian", "Punjabi", "Pushto", "Sanskrit", "Serbian", "Shona", "Sindhi", "Sinhala", "Sinhalese", "Slovak", "Slovenian", "Somali", "Spanish", "Sundanese", "Swahili", "Tagalog", "Tajik", "Tamil", "Tatar", "Telugu", "Tibetan", "Turkmen", "Urdu", "Uzbek", "Valencian", "Yoruba"],
        index=0
    )
    
    # Word timestamps option
    word_timestamps = st.checkbox(
        "Include word-level timestamps",
        value=False,
        help="Extract timestamps for each word (slower but more detailed)"
    )

with tab1:
    # Note about large files
    st.info("ℹ️ **Large File Support**: Files up to 2GB are supported. Files >200MB will be automatically split into chunks and processed in parallel. If you see a 200MB limit, please restart the Streamlit app.")
    
    # File upload - increase size limit for large files
    uploaded_file = st.file_uploader(
        "Upload Audio or Video File",
        type=["mp3", "wav", "flac", "m4a", "ogg", "wma", "aac", "mov", "mp4", "avi", "mkv", "webm", "m4v"],
        help="Supported formats: Audio (MP3, WAV, FLAC, M4A, OGG, WMA, AAC) and Video (MOV, MP4, AVI, MKV, WEBM, M4V). Large files (>200MB) will be automatically split into chunks and processed in parallel. Maximum upload size: 2GB (requires app restart if you see 200MB limit)."
    )

    if uploaded_file is not None:
        # Display file info
        file_details = {
            "Filename": uploaded_file.name,
            "FileType": uploaded_file.type,
            "FileSize": f"{uploaded_file.size / 1024 / 1024:.2f} MB"
        }
        
        st.info(f"📁 **{file_details['Filename']}** ({file_details['FileSize']})")
        
        # Process button
        if st.button("🚀 Transcribe Audio", type="primary", use_container_width=True):
            # Create containers for progress updates
            progress_container = st.container()
            log_container = st.container()
            
            try:
                # Step 1: Save uploaded file
                with progress_container:
                    st.toast("📁 Saving uploaded file...", icon="📁")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_file_path = tmp_file.name
                    st.toast("✅ File saved successfully", icon="✅")
                
                # Step 2: Load model
                with progress_container:
                    st.toast("🔄 Loading model... This may take a moment on first use", icon="🔄")
                    progress_bar = st.progress(0, text="Loading model...")
                    status_text = st.empty()
                    
                    status_text.info("📦 Loading Whisper model (this may download the model on first use)...")
                    model = whisper.load_model(selected_model)
                    
                    progress_bar.progress(100, text="Model loaded!")
                    status_text.success("✅ Model loaded successfully!")
                    st.toast("✅ Model loaded!", icon="✅")
                
                # Step 3: Prepare transcription options
                transcribe_options = {
                    "task": task,
                    "word_timestamps": word_timestamps,
                    "verbose": True  # Enable verbose to get progress updates
                }
                
                if language != "Auto-detect":
                    transcribe_options["language"] = language
                    st.toast(f"🌍 Language set to: {language}", icon="🌍")
                else:
                    st.toast("🔍 Auto-detecting language...", icon="🔍")
                
                # Step 4: Check file size and decide on processing method
                file_size_mb = uploaded_file.size / (1024 * 1024)
                use_chunking = file_size_mb > CHUNK_SIZE_MB
                
                if use_chunking:
                    # Large file: Use chunking and parallel processing
                    with progress_container:
                        st.toast("📦 Large file detected! Splitting into chunks...", icon="📦")
                        transcription_progress = st.progress(0, text="Preparing chunks...")
                        transcription_status = st.empty()
                        transcription_status.info(f"📦 File is {file_size_mb:.2f} MB. Splitting into chunks for parallel processing...")
                        
                        # Split file into chunks
                        chunks, chunk_dir = split_file_into_chunks(tmp_file_path, max_workers=MAX_WORKERS)
                        num_chunks = len(chunks)
                        
                        if num_chunks == 0:
                            raise ValueError("Failed to create valid chunks from the file. The file may be corrupted or contain no audio.")
                        
                        # Calculate actual thread usage (min of workers and chunks)
                        actual_threads = min(MAX_WORKERS, num_chunks)
                        
                        transcription_progress.progress(10, text=f"Created {num_chunks} valid chunks. Starting parallel transcription...")
                        transcription_status.info(f"🔄 Processing {num_chunks} chunks in parallel using {actual_threads} threads (pool size: {MAX_WORKERS})...")
                        
                        # Show chunk coverage info
                        if chunks:
                            total_chunk_duration = sum(chunk.get("duration", 0) for chunk in chunks)
                            file_duration = get_file_duration(tmp_file_path)
                            if file_duration:
                                coverage_pct = (total_chunk_duration / file_duration) * 100
                                st.info(f"📊 Chunk coverage: {total_chunk_duration:.1f}s / {file_duration:.1f}s ({coverage_pct:.1f}%) - All chunks will be processed and reconstituted")
                        
                        # Create individual progress indicators for each chunk
                        chunk_progress_container = st.container()
                        with chunk_progress_container:
                            st.subheader("📊 Individual Chunk Progress")
                            chunk_progress_bars = {}
                            chunk_status_texts = {}
                            
                            # Initialize progress bars for each chunk
                            for chunk in chunks:
                                chunk_idx = chunk["index"]
                                chunk_key = f"chunk_{chunk_idx}"
                                
                                col1, col2 = st.columns([1, 4])
                                with col1:
                                    st.write(f"**Chunk {chunk_idx + 1}**")
                                    st.caption(f"{chunk['start_time']:.1f}s - {chunk['end_time']:.1f}s")
                                
                                with col2:
                                    chunk_progress_bars[chunk_key] = st.progress(0, text="Waiting to start...")
                                    chunk_status_texts[chunk_key] = st.empty()
                                    chunk_status_texts[chunk_key].info("⏳ Queued for processing...")
                        
                        # Process chunks in parallel
                        chunk_results = []
                        completed_chunks = 0
                        successful_chunks = 0
                        failed_chunks = []
                        
                        # Thread-safe progress tracking and model access
                        progress_lock = threading.Lock()
                        model_lock = threading.Lock()  # Lock for model access (Whisper models are not thread-safe)
                        progress_updates = {}  # Store progress updates to apply in main thread
                        
                        def update_chunk_progress(chunk_idx, status, progress_pct=0):
                            """Update progress for a specific chunk (thread-safe)"""
                            chunk_key = f"chunk_{chunk_idx}"
                            with progress_lock:
                                progress_updates[chunk_key] = {
                                    "status": status,
                                    "progress": progress_pct,
                                    "message": f"🔄 {status}"
                                }
                        
                        def transcribe_chunk_with_progress(chunk_info, model, transcribe_options):
                            """Transcribe chunk with progress updates"""
                            chunk_idx = chunk_info["index"]
                            chunk_key = f"chunk_{chunk_idx}"
                            
                            try:
                                update_chunk_progress(chunk_idx, "Starting transcription...", 10)
                                
                                # Validate chunk first
                                if not os.path.exists(chunk_info["path"]):
                                    update_chunk_progress(chunk_idx, "❌ File not found", 0)
                                    return {
                                        "chunk_index": chunk_idx,
                                        "start_time": chunk_info["start_time"],
                                        "error": "Chunk file does not exist",
                                        "success": False
                                    }
                                
                                file_size = os.path.getsize(chunk_info["path"])
                                if file_size < 1024:
                                    update_chunk_progress(chunk_idx, "❌ Empty chunk", 0)
                                    return {
                                        "chunk_index": chunk_idx,
                                        "start_time": chunk_info["start_time"],
                                        "error": "Chunk file is too small (likely empty)",
                                        "success": False
                                    }
                                
                                update_chunk_progress(chunk_idx, "Loading audio...", 20)
                                
                                # Load and validate audio
                                try:
                                    audio = whisper.load_audio(chunk_info["path"])
                                    if len(audio) == 0 or audio.size == 0:
                                        update_chunk_progress(chunk_idx, "❌ No audio data", 0)
                                        return {
                                            "chunk_index": chunk_idx,
                                            "start_time": chunk_info["start_time"],
                                            "error": "Chunk contains no audio data",
                                            "success": False
                                        }
                                except Exception as e:
                                    update_chunk_progress(chunk_idx, f"❌ Audio load failed: {str(e)[:50]}", 0)
                                    return {
                                        "chunk_index": chunk_idx,
                                        "start_time": chunk_info["start_time"],
                                        "error": f"Failed to load audio: {str(e)}",
                                        "success": False
                                    }
                                
                                update_chunk_progress(chunk_idx, "Transcribing...", 40)
                                
                                # Transcribe with progress tracking and better error handling
                                # Use model lock to ensure thread-safe access (Whisper models are not thread-safe)
                                try:
                                    with model_lock:
                                        result = model.transcribe(chunk_info["path"], **transcribe_options)
                                except Exception as transcribe_error:
                                    # Catch transcription errors and provide better error messages
                                    error_msg = str(transcribe_error)
                                    # Filter out model internals from error messages
                                    if "Linear(" in error_msg or "in_features" in error_msg:
                                        error_msg = "Model error during transcription (possibly corrupted audio or model issue)"
                                    elif "CUDA" in error_msg or "cuda" in error_msg:
                                        error_msg = "GPU/CUDA error during transcription"
                                    elif "out of memory" in error_msg.lower():
                                        error_msg = "Out of memory during transcription"
                                    else:
                                        # Keep original error but truncate if too long
                                        error_msg = error_msg[:200] if len(error_msg) > 200 else error_msg
                                    
                                    update_chunk_progress(chunk_idx, f"❌ Transcription failed", 0)
                                    return {
                                        "chunk_index": chunk_idx,
                                        "start_time": chunk_info["start_time"],
                                        "error": error_msg,
                                        "success": False
                                    }
                                
                                update_chunk_progress(chunk_idx, "Processing segments...", 80)
                                
                                # Adjust timestamps to account for chunk offset
                                if result.get("segments"):
                                    for segment in result["segments"]:
                                        segment["start"] += chunk_info["start_time"]
                                        segment["end"] += chunk_info["start_time"]
                                        if segment.get("words"):
                                            for word in segment["words"]:
                                                word["start"] += chunk_info["start_time"]
                                                word["end"] += chunk_info["start_time"]
                                
                                num_segments = len(result.get('segments', []))
                                update_chunk_progress(chunk_idx, f"✅ Complete ({num_segments} segments)", 100)
                                
                                return {
                                    "chunk_index": chunk_idx,
                                    "start_time": chunk_info["start_time"],
                                    "result": result,
                                    "success": True,
                                    "num_segments": num_segments
                                }
                            except Exception as e:
                                error_msg = str(e)
                                if "reshape" in error_msg and "0 elements" in error_msg:
                                    error_msg = "Chunk contains no audio data or is empty"
                                update_chunk_progress(chunk_idx, f"❌ Error: {error_msg[:50]}", 0)
                                return {
                                    "chunk_index": chunk_idx,
                                    "start_time": chunk_info["start_time"],
                                    "error": error_msg,
                                    "success": False
                                }
                        
                        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                            # Submit all chunks with progress tracking
                            future_to_chunk = {
                                executor.submit(transcribe_chunk_with_progress, chunk, model, transcribe_options): chunk
                                for chunk in chunks
                            }
                            
                            # Process completed chunks as they finish and update UI
                            for future in as_completed(future_to_chunk):
                                # Apply any pending progress updates
                                with progress_lock:
                                    for chunk_key, update_info in progress_updates.items():
                                        if chunk_key in chunk_progress_bars:
                                            chunk_progress_bars[chunk_key].progress(
                                                update_info["progress"] / 100, 
                                                text=update_info["status"]
                                            )
                                            chunk_status_texts[chunk_key].info(update_info["message"])
                                    progress_updates.clear()
                                
                                chunk_result = future.result()
                                chunk_results.append(chunk_result)
                                completed_chunks += 1
                                
                                chunk_idx = chunk_result["chunk_index"]
                                chunk_key = f"chunk_{chunk_idx}"
                                
                                if chunk_result["success"]:
                                    successful_chunks += 1
                                    num_segments = chunk_result.get("num_segments", 0)
                                    chunk_progress_bars[chunk_key].progress(100, text=f"✅ Complete ({num_segments} segments)")
                                    chunk_status_texts[chunk_key].success(f"✅ Completed: {num_segments} segments")
                                    st.toast(f"✅ Chunk {chunk_idx + 1}/{num_chunks} completed", icon="✅")
                                else:
                                    failed_chunks.append({
                                        "index": chunk_idx + 1,
                                        "error": chunk_result.get('error', 'Unknown error')
                                    })
                                    chunk_progress_bars[chunk_key].progress(0, text=f"❌ Failed")
                                    chunk_status_texts[chunk_key].error(f"❌ Failed: {chunk_result.get('error', 'Unknown error')[:100]}")
                                    # Don't show warning for empty chunks (expected for silent segments)
                                    if "empty" not in chunk_result.get('error', '').lower() and "no audio" not in chunk_result.get('error', '').lower():
                                        st.warning(f"⚠️ Chunk {chunk_idx + 1} skipped: {chunk_result.get('error', 'Unknown error')}")
                                
                                # Update overall progress
                                progress_pct = 10 + int((completed_chunks / num_chunks) * 85)
                                transcription_progress.progress(
                                    progress_pct / 100,
                                    text=f"Processing chunks... {completed_chunks}/{num_chunks} completed ({successful_chunks} successful)"
                                )
                        
                        # Check if we have any successful chunks
                        if successful_chunks == 0:
                            raise ValueError(f"All {num_chunks} chunks failed to transcribe. Please check the file format and try again.")
                        
                        # Calculate success rate
                        success_rate = (successful_chunks / num_chunks) * 100
                        
                        # Show summary of failed chunks if any
                        if failed_chunks:
                            failed_summary = "\n".join([f"  - Chunk {f['index']}: {f['error']}" for f in failed_chunks])
                            
                            # Determine severity based on success rate
                            if success_rate < 50:
                                # Less than 50% success - this is a problem
                                transcription_status.error(f"❌ **WARNING**: Only {successful_chunks}/{num_chunks} chunks succeeded ({success_rate:.1f}% success rate). Transcription may be incomplete!\n\n**Failed chunks:**\n{failed_summary}\n\n**Recommendation**: Check the file for corruption or try a different model size.")
                            elif success_rate < 80:
                                # 50-80% success - partial success
                                transcription_status.warning(f"⚠️ **Partial Success**: {successful_chunks}/{num_chunks} chunks succeeded ({success_rate:.1f}% success rate). Some content may be missing.\n\n**Failed chunks:**\n{failed_summary}\n\nContinuing with {successful_chunks} successful chunk(s)...")
                            else:
                                # 80%+ success - mostly successful
                                transcription_status.warning(f"⚠️ {len(failed_chunks)} chunk(s) failed:\n{failed_summary}\n\nContinuing with {successful_chunks} successful chunk(s)...")
                        
                        # Combine results
                        transcription_progress.progress(95, text="Combining results...")
                        transcription_status.info(f"🔗 Combining transcription results from {successful_chunks} successful chunk(s)...")
                        result = combine_chunk_results(chunk_results, chunks)
                        
                        # Verify completeness
                        if result.get("_metadata"):
                            missing = result["_metadata"].get("missing_chunks", [])
                            if missing:
                                st.warning(f"⚠️ Note: {len(missing)} chunk(s) were skipped: {missing}")
                        
                        # Show summary with success rate
                        total_segments = len(result.get("segments", []))
                        total_text_length = len(result.get("text", ""))
                        
                        # Determine final status based on success rate
                        if success_rate >= 80:
                            transcription_status.success(f"✅ **Transcription completed!** {successful_chunks}/{num_chunks} chunks processed ({success_rate:.1f}% success), {total_segments} segments, {total_text_length:,} characters")
                        elif success_rate >= 50:
                            transcription_status.warning(f"⚠️ **Partial transcription completed.** {successful_chunks}/{num_chunks} chunks processed ({success_rate:.1f}% success), {total_segments} segments, {total_text_length:,} characters. Some content may be missing.")
                        else:
                            transcription_status.error(f"❌ **Incomplete transcription.** Only {successful_chunks}/{num_chunks} chunks processed ({success_rate:.1f}% success), {total_segments} segments, {total_text_length:,} characters. Significant content may be missing.")
                        
                        # Clean up chunk files
                        import shutil
                        try:
                            shutil.rmtree(chunk_dir)
                        except:
                            pass
                        
                        transcription_progress.progress(100, text="Transcription complete!")
                        
                        # Final toast message based on success rate
                        if success_rate >= 80:
                            st.toast("✅ Transcription completed successfully!", icon="✅")
                        elif success_rate >= 50:
                            st.toast("⚠️ Partial transcription completed - some chunks failed", icon="⚠️")
                        else:
                            st.toast("❌ Incomplete transcription - many chunks failed", icon="❌")
                else:
                    # Small file: Use standard processing
                    with progress_container:
                        st.toast("🎤 Starting transcription...", icon="🎤")
                        transcription_progress = st.progress(0, text="Initializing transcription...")
                        transcription_status = st.empty()
                        transcription_status.info("🎤 Transcribing audio... This may take a while depending on file size.")
                        
                        # Create a collapsible log section
                        with st.expander("📋 View Transcription Logs (Real-time Progress)", expanded=True):
                            log_placeholder = st.empty()
                            log_output = []
                            
                            # Capture stdout to show transcription progress
                            class ProgressCapture:
                                def __init__(self):
                                    self.lines = []
                                    self.placeholder = log_placeholder
                                    self.segment_count = 0
                                
                                def write(self, text):
                                    if text.strip():
                                        # Parse segment lines (format: [00:00.000 --> 00:11.000] text)
                                        line = text.strip()
                                        self.lines.append(line)
                                        
                                        # Count segments
                                        if "-->" in line and "]" in line:
                                            self.segment_count += 1
                                            # Update progress estimate (rough estimate: 5% per segment)
                                            estimated_progress = min(95, self.segment_count * 5)
                                            transcription_progress.progress(
                                                estimated_progress / 100, 
                                                text=f"Transcribing... ({self.segment_count} segments processed)"
                                            )
                                        
                                        # Update log display (show last 15 lines)
                                        display_lines = self.lines[-15:]
                                        self.placeholder.code("\n".join(display_lines), language="text")
                                
                                def flush(self):
                                    pass
                            
                            progress_capture = ProgressCapture()
                            
                            # Transcribe with verbose output captured
                            with redirect_stdout(progress_capture):
                                with redirect_stderr(progress_capture):
                                    result = model.transcribe(tmp_file_path, **transcribe_options)
                            
                            # Final log update
                            if progress_capture.lines:
                                log_placeholder.code("\n".join(progress_capture.lines[-20:]), language="text")
                    
                        transcription_progress.progress(100, text="Transcription complete!")
                        transcription_status.success(f"✅ Transcription completed! Processed {len(result.get('segments', []))} segments.")
                        st.toast("✅ Transcription complete!", icon="✅")
                
                # Step 5: Clean up
                os.unlink(tmp_file_path)
                
                # Step 6: Save transcription to storage
                transcription_metadata = {
                    "id": str(int(time.time())),
                    "filename": uploaded_file.name,
                    "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "language": result.get("language", "Unknown"),
                    "duration": result.get("duration", 0),
                    "segments_count": len(result.get("segments", [])),
                    "model": selected_model,
                    "task": task,
                    "text": result["text"],
                    "full_result": result  # Store full result for JSON download
                }
                save_transcription(transcription_metadata)
                st.toast("💾 Transcription saved to storage", icon="💾")
                
                # Display results
                st.balloons()  # Celebration animation
                st.success("🎉 Transcription completed successfully!")
                
                # Main transcription text
                st.subheader("📝 Transcription")
                st.text_area(
                    "Transcribed Text",
                    value=result["text"],
                    height=200,
                    label_visibility="collapsed",
                    key="main_transcription_text_area"
                )
                
                # Additional information
                with st.expander("ℹ️ Additional Information"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.metric("Detected Language", result.get("language", "N/A"))
                        st.metric("Duration", f"{result.get('duration', 0):.2f} seconds")
                    
                    with col2:
                        st.metric("Segments", len(result.get("segments", [])))
                        if result.get("language"):
                            # Get language probability if available
                            st.metric("Task", task.capitalize())
                
                # Segments with timestamps
                if result.get("segments"):
                    st.subheader("⏱️ Segments with Timestamps")
                    
                    for i, segment in enumerate(result["segments"]):
                        start_time = segment.get("start", 0)
                        end_time = segment.get("end", 0)
                        
                        # Format time
                        def format_time(seconds):
                            hours = int(seconds // 3600)
                            minutes = int((seconds % 3600) // 60)
                            secs = int(seconds % 60)
                            if hours > 0:
                                return f"{hours:02d}:{minutes:02d}:{secs:02d}"
                            return f"{minutes:02d}:{secs:02d}"
                        
                        st.markdown(f"**{format_time(start_time)} → {format_time(end_time)}**")
                        st.write(segment.get("text", ""))
                        
                        if word_timestamps and segment.get("words"):
                            st.caption("Words: " + " | ".join([
                                f"{w.get('word', '')} ({w.get('start', 0):.2f}s)"
                                for w in segment.get("words", [])[:10]  # Show first 10 words
                            ]))
                        
                        if i < len(result["segments"]) - 1:
                            st.divider()
                
                # Download options
                st.subheader("💾 Download Results")
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    # Create TXT with timestamps
                    txt_content = ""
                    if result.get("segments"):
                        # Add header
                        txt_content += f"Transcription: {uploaded_file.name}\n"
                        txt_content += f"Language: {result.get('language', 'Unknown')}\n"
                        txt_content += f"Duration: {result.get('duration', 0):.2f} seconds\n"
                        txt_content += f"Date: {transcription_metadata['date']}\n"
                        txt_content += "=" * 50 + "\n\n"
                        
                        # Add segments with timestamps
                        txt_content += "TRANSCRIPTION WITH TIMESTAMPS:\n"
                        txt_content += "-" * 50 + "\n\n"
                        for segment in result["segments"]:
                            start = segment.get('start', 0)
                            end = segment.get('end', 0)
                            text = segment.get('text', '').strip()
                            
                            # Format timestamp
                            def format_time(seconds):
                                hours = int(seconds // 3600)
                                minutes = int((seconds % 3600) // 60)
                                secs = int(seconds % 60)
                                millis = int((seconds % 1) * 1000)
                                if hours > 0:
                                    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
                                return f"{minutes:02d}:{secs:02d}.{millis:03d}"
                            
                            txt_content += f"[{format_time(start)} --> {format_time(end)}] {text}\n\n"
                        
                        txt_content += "\n" + "=" * 50 + "\n\n"
                        txt_content += "PLAIN TEXT (NO TIMESTAMPS):\n"
                        txt_content += "-" * 50 + "\n\n"
                    
                    # Add plain text
                    txt_content += result["text"]
                    
                    st.download_button(
                        label="📄 Download as TXT",
                        data=txt_content,
                        file_name=f"{Path(uploaded_file.name).stem}_transcription.txt",
                        mime="text/plain"
                    )
                
                with col2:
                    # Create DOCX file with timestamps
                    docx_buffer = create_docx(
                        result["text"],
                        uploaded_file.name,
                        {
                            "filename": uploaded_file.name,
                            "language": result.get("language", "Unknown"),
                            "duration": result.get("duration", 0),
                            "date": transcription_metadata["date"]
                        },
                        segments=result.get("segments")
                    )
                    st.download_button(
                        label="📝 Download as DOCX",
                        data=docx_buffer.getvalue(),
                        file_name=f"{Path(uploaded_file.name).stem}_transcription.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                with col3:
                    # Create JSON output
                    json_output = json.dumps(result, indent=2, ensure_ascii=False)
                    st.download_button(
                        label="📊 Download as JSON",
                        data=json_output,
                        file_name=f"{Path(uploaded_file.name).stem}_transcription.json",
                        mime="application/json"
                    )
                
                with col4:
                    # Create SRT subtitle format
                    if result.get("segments"):
                        srt_content = ""
                        for i, segment in enumerate(result["segments"], 1):
                            start = segment.get("start", 0)
                            end = segment.get("end", 0)
                            
                            def format_srt_time(seconds):
                                hours = int(seconds // 3600)
                                minutes = int((seconds % 3600) // 60)
                                secs = int(seconds % 60)
                                millis = int((seconds % 1) * 1000)
                                return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
                            
                            srt_content += f"{i}\n"
                            srt_content += f"{format_srt_time(start)} --> {format_srt_time(end)}\n"
                            srt_content += f"{segment.get('text', '')}\n\n"
                        
                        st.download_button(
                            label="📺 Download as SRT",
                            data=srt_content,
                            file_name=f"{Path(uploaded_file.name).stem}_subtitles.srt",
                            mime="text/plain"
                        )
            
            except Exception as e:
                st.error(f"❌ Error during transcription: {str(e)}")
                st.exception(e)

with tab2:
    st.header("📚 Stored Transcriptions")
    
    # Tabs for active and archived transcriptions
    view_tab1, view_tab2 = st.tabs(["📋 Active", "📦 Archived"])
    
    with view_tab1:
        transcriptions = load_transcriptions()
        
        if not transcriptions:
            st.info("📭 No stored transcriptions yet. Complete a transcription to see it here.")
        else:
            # Sort by date (newest first)
            transcriptions.sort(key=lambda x: x.get("date", ""), reverse=True)
            
            col_metric, col_search = st.columns([1, 3])
            with col_metric:
                st.metric("Total Transcriptions", len(transcriptions))
            with col_search:
                search_term = st.text_input("🔍 Search transcriptions", placeholder="Search by filename or date...", key="search_active")
            
            # Filter transcriptions
            filtered_transcriptions = transcriptions
            if search_term:
                filtered_transcriptions = [
                    t for t in transcriptions
                    if search_term.lower() in t.get("filename", "").lower() 
                    or search_term.lower() in t.get("date", "").lower()
                    or search_term.lower() in t.get("text", "").lower()
                ]
            
            if not filtered_transcriptions:
                st.warning("No transcriptions found matching your search.")
            else:
                for idx, trans in enumerate(filtered_transcriptions):
                    trans_id = trans.get('id', 'unknown')
                    
                    # Transcription viewer with timestamp navigation
                    with st.expander(
                        f"📄 {trans.get('filename', 'Unknown')} - {trans.get('date', 'Unknown date')}",
                        expanded=False
                    ):
                        # Metadata
                        col1, col2 = st.columns([2, 1])
                        
                        with col1:
                            st.write(f"**Language:** {trans.get('language', 'Unknown')}")
                            st.write(f"**Duration:** {trans.get('duration', 0):.2f} seconds")
                            st.write(f"**Segments:** {trans.get('segments_count', 0)}")
                            st.write(f"**Model:** {trans.get('model', 'Unknown')}")
                            st.write(f"**Task:** {trans.get('task', 'transcribe').capitalize()}")
                        
                        with col2:
                            st.write(f"**Date:** {trans.get('date', 'Unknown')}")
                            st.write(f"**ID:** {trans.get('id', 'Unknown')}")
                        
                        # Timestamp-based viewer
                        full_result = trans.get('full_result', {})
                        segments = full_result.get("segments", [])
                        
                        if segments:
                            st.subheader("🎬 Transcription Viewer with Timestamps")
                            
                            # Initialize selectbox key in session state if not exists
                            selectbox_key = f"segment_select_{trans_id}"
                            if selectbox_key not in st.session_state:
                                st.session_state[selectbox_key] = 0
                            
                            # Handle navigation button clicks
                            nav_col1, nav_col2, nav_col3, nav_col4 = st.columns(4)
                            with nav_col1:
                                if st.button("⏮️ First", key=f"first_{trans_id}"):
                                    st.session_state[selectbox_key] = 0
                                    st.rerun()
                            with nav_col2:
                                if st.button("◀️ Previous", key=f"prev_{trans_id}"):
                                    current_idx = st.session_state.get(selectbox_key, 0)
                                    if current_idx > 0:
                                        st.session_state[selectbox_key] = current_idx - 1
                                        st.rerun()
                            with nav_col3:
                                if st.button("Next ▶️", key=f"next_{trans_id}"):
                                    current_idx = st.session_state.get(selectbox_key, 0)
                                    if current_idx < len(segments) - 1:
                                        st.session_state[selectbox_key] = current_idx + 1
                                        st.rerun()
                            with nav_col4:
                                if st.button("⏭️ Last", key=f"last_{trans_id}"):
                                    st.session_state[selectbox_key] = len(segments) - 1
                                    st.rerun()
                            
                            # Get current segment index from session state
                            current_segment_idx = st.session_state.get(selectbox_key, 0)
                            
                            # Ensure index is within bounds
                            if current_segment_idx >= len(segments):
                                current_segment_idx = len(segments) - 1
                                st.session_state[selectbox_key] = current_segment_idx
                            if current_segment_idx < 0:
                                current_segment_idx = 0
                                st.session_state[selectbox_key] = 0
                            
                            # Segment navigation dropdown
                            segment_options = [f"[{format_timestamp(s.get('start', 0))} → {format_timestamp(s.get('end', 0))}] {s.get('text', '')[:50]}..." 
                                             for s in segments]
                            
                            # Selectbox maintains its own state via key
                            selected_segment_idx = st.selectbox(
                                "Navigate by timestamp:",
                                range(len(segments)),
                                index=current_segment_idx,
                                format_func=lambda x: segment_options[x] if x < len(segment_options) else f"Segment {x+1}",
                                key=selectbox_key
                            )
                            
                            # Display selected segment
                            if selected_segment_idx < len(segments):
                                selected_segment = segments[selected_segment_idx]
                                start_time = selected_segment.get('start', 0)
                                end_time = selected_segment.get('end', 0)
                                segment_text = selected_segment.get('text', '').strip()
                                
                                st.markdown(f"**Timestamp:** `{format_timestamp(start_time)}` → `{format_timestamp(end_time)}`")
                                st.text_area(
                                    "Segment Text",
                                    value=segment_text,
                                    height=100,
                                    disabled=True,
                                    key=f"segment_text_{trans_id}_{selected_segment_idx}"
                                )
                            
                            # Full transcription view
                            st.subheader("📝 Full Transcription")
                            full_text = trans.get('text', '')
                            st.text_area(
                                "Complete Transcription",
                                value=full_text,
                                height=300,
                                disabled=True,
                                key=f"full_text_{trans_id}"
                            )
                        else:
                            # Fallback if no segments
                            st.text_area(
                                "Transcription",
                                value=trans.get('text', ''),
                                height=200,
                                disabled=True,
                                key=f"text_{trans_id}"
                            )
                    
                    # Download buttons
                    st.subheader("💾 Download")
                    dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
                    
                    with dl_col1:
                        # Create TXT with timestamps for stored transcription
                        txt_content = ""
                        full_result = trans.get('full_result', {})
                        if full_result.get("segments"):
                            # Add header
                            txt_content += f"Transcription: {trans.get('filename', 'Unknown')}\n"
                            txt_content += f"Language: {trans.get('language', 'Unknown')}\n"
                            txt_content += f"Duration: {trans.get('duration', 0):.2f} seconds\n"
                            txt_content += f"Date: {trans.get('date', 'Unknown')}\n"
                            txt_content += "=" * 50 + "\n\n"
                            
                            # Add segments with timestamps
                            txt_content += "TRANSCRIPTION WITH TIMESTAMPS:\n"
                            txt_content += "-" * 50 + "\n\n"
                            for segment in full_result["segments"]:
                                start = segment.get('start', 0)
                                end = segment.get('end', 0)
                                text = segment.get('text', '').strip()
                                
                                # Format timestamp
                                def format_time(seconds):
                                    hours = int(seconds // 3600)
                                    minutes = int((seconds % 3600) // 60)
                                    secs = int(seconds % 60)
                                    millis = int((seconds % 1) * 1000)
                                    if hours > 0:
                                        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
                                    return f"{minutes:02d}:{secs:02d}.{millis:03d}"
                                
                                txt_content += f"[{format_time(start)} --> {format_time(end)}] {text}\n\n"
                            
                            txt_content += "\n" + "=" * 50 + "\n\n"
                            txt_content += "PLAIN TEXT (NO TIMESTAMPS):\n"
                            txt_content += "-" * 50 + "\n\n"
                        
                        # Add plain text
                        txt_content += trans.get('text', '')
                        
                        st.download_button(
                            label="📄 TXT",
                            data=txt_content,
                            file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.txt",
                            mime="text/plain",
                            key=f"txt_{trans.get('id')}"
                        )
                    
                    with dl_col2:
                        # Create DOCX for stored transcription
                        full_result = trans.get('full_result', {})
                        docx_buffer = create_docx(
                            trans.get('text', ''),
                            trans.get('filename', 'transcription'),
                            {
                                "filename": trans.get('filename', 'Unknown'),
                                "language": trans.get('language', 'Unknown'),
                                "duration": trans.get('duration', 0),
                                "date": trans.get('date', 'Unknown')
                            },
                            segments=full_result.get("segments")
                        )
                        st.download_button(
                            label="📝 DOCX",
                            data=docx_buffer.getvalue(),
                            file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"docx_{trans.get('id')}"
                        )
                    
                    with dl_col3:
                        if trans.get('full_result'):
                            json_output = json.dumps(trans['full_result'], indent=2, ensure_ascii=False)
                            st.download_button(
                                label="📊 JSON",
                                data=json_output,
                                file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.json",
                                mime="application/json",
                                key=f"json_{trans.get('id')}"
                            )
                    
                    with dl_col4:
                        # Archive and Purge buttons
                        action_col1, action_col2 = st.columns(2)
                        with action_col1:
                            if st.button("📦 Archive", key=f"archive_{trans.get('id')}", help="Move to archive"):
                                archive_transcription(trans.get('id'))
                                st.success("✅ Transcription archived!")
                                st.rerun()
                        with action_col2:
                            if st.button("🗑️ Purge", key=f"purge_{trans.get('id')}", help="Permanently delete", type="secondary"):
                                purge_transcription(trans.get('id'))
                                st.success("✅ Transcription purged!")
                                st.rerun()
                    
                    if idx < len(filtered_transcriptions) - 1:
                        st.divider()
    
    with view_tab2:
        archived = load_archived_transcriptions()
        
        if not archived:
            st.info("📦 No archived transcriptions yet.")
        else:
            # Sort by archived date (newest first)
            archived.sort(key=lambda x: x.get("archived_date", ""), reverse=True)
            
            col_metric, col_search = st.columns([1, 3])
            with col_metric:
                st.metric("Archived Transcriptions", len(archived))
            with col_search:
                search_term = st.text_input("🔍 Search archived", placeholder="Search by filename or date...", key="search_archived")
            
            # Filter archived
            filtered_archived = archived
            if search_term:
                filtered_archived = [
                    t for t in archived
                    if search_term.lower() in t.get("filename", "").lower() 
                    or search_term.lower() in t.get("date", "").lower()
                    or search_term.lower() in t.get("text", "").lower()
                ]
            
            if not filtered_archived:
                st.warning("No archived transcriptions found matching your search.")
            else:
                for idx, trans in enumerate(filtered_archived):
                    trans_id = trans.get('id', 'unknown')
                    
                    with st.expander(
                        f"📦 {trans.get('filename', 'Unknown')} - Archived: {trans.get('archived_date', 'Unknown date')}",
                        expanded=False
                    ):
                        # Metadata
                        col1, col2 = st.columns([2, 1])
                        
                        with col1:
                            st.write(f"**Language:** {trans.get('language', 'Unknown')}")
                            st.write(f"**Duration:** {trans.get('duration', 0):.2f} seconds")
                            st.write(f"**Segments:** {trans.get('segments_count', 0)}")
                            st.write(f"**Model:** {trans.get('model', 'Unknown')}")
                            st.write(f"**Task:** {trans.get('task', 'transcribe').capitalize()}")
                        
                        with col2:
                            st.write(f"**Original Date:** {trans.get('date', 'Unknown')}")
                            st.write(f"**Archived Date:** {trans.get('archived_date', 'Unknown')}")
                            st.write(f"**ID:** {trans.get('id', 'Unknown')}")
                        
                        # Timestamp-based viewer (same as active)
                        full_result = trans.get('full_result', {})
                        segments = full_result.get("segments", [])
                        
                        if segments:
                            st.subheader("🎬 Transcription Viewer with Timestamps")
                            
                            # Initialize selectbox key in session state if not exists
                            selectbox_key = f"archived_segment_select_{trans_id}"
                            if selectbox_key not in st.session_state:
                                st.session_state[selectbox_key] = 0
                            
                            # Handle navigation button clicks
                            nav_col1, nav_col2, nav_col3, nav_col4 = st.columns(4)
                            with nav_col1:
                                if st.button("⏮️ First", key=f"archived_first_{trans_id}"):
                                    st.session_state[selectbox_key] = 0
                                    st.rerun()
                            with nav_col2:
                                if st.button("◀️ Previous", key=f"archived_prev_{trans_id}"):
                                    current_idx = st.session_state.get(selectbox_key, 0)
                                    if current_idx > 0:
                                        st.session_state[selectbox_key] = current_idx - 1
                                        st.rerun()
                            with nav_col3:
                                if st.button("Next ▶️", key=f"archived_next_{trans_id}"):
                                    current_idx = st.session_state.get(selectbox_key, 0)
                                    if current_idx < len(segments) - 1:
                                        st.session_state[selectbox_key] = current_idx + 1
                                        st.rerun()
                            with nav_col4:
                                if st.button("⏭️ Last", key=f"archived_last_{trans_id}"):
                                    st.session_state[selectbox_key] = len(segments) - 1
                                    st.rerun()
                            
                            # Get current segment index from session state
                            current_segment_idx = st.session_state.get(selectbox_key, 0)
                            
                            # Ensure index is within bounds
                            if current_segment_idx >= len(segments):
                                current_segment_idx = len(segments) - 1
                                st.session_state[selectbox_key] = current_segment_idx
                            if current_segment_idx < 0:
                                current_segment_idx = 0
                                st.session_state[selectbox_key] = 0
                            
                            # Segment navigation dropdown
                            segment_options = [f"[{format_timestamp(s.get('start', 0))} → {format_timestamp(s.get('end', 0))}] {s.get('text', '')[:50]}..." 
                                             for s in segments]
                            
                            # Selectbox maintains its own state via key
                            selected_segment_idx = st.selectbox(
                                "Navigate by timestamp:",
                                range(len(segments)),
                                index=current_segment_idx,
                                format_func=lambda x: segment_options[x] if x < len(segment_options) else f"Segment {x+1}",
                                key=selectbox_key
                            )
                            
                            # Display selected segment
                            if selected_segment_idx < len(segments):
                                selected_segment = segments[selected_segment_idx]
                                start_time = selected_segment.get('start', 0)
                                end_time = selected_segment.get('end', 0)
                                segment_text = selected_segment.get('text', '').strip()
                                
                                st.markdown(f"**Timestamp:** `{format_timestamp(start_time)}` → `{format_timestamp(end_time)}`")
                                st.text_area(
                                    "Segment Text",
                                    value=segment_text,
                                    height=100,
                                    disabled=True,
                                    key=f"archived_segment_text_{trans_id}_{selected_segment_idx}"
                                )
                            
                            # Full transcription view
                            st.subheader("📝 Full Transcription")
                            full_text = trans.get('text', '')
                            st.text_area(
                                "Complete Transcription",
                                value=full_text,
                                height=300,
                                disabled=True,
                                key=f"archived_full_text_{trans_id}"
                            )
                        else:
                            st.text_area(
                                "Transcription",
                                value=trans.get('text', ''),
                                height=200,
                                disabled=True,
                                key=f"archived_text_{trans_id}"
                            )
                        
                        # Download buttons (same as active)
                        st.subheader("💾 Download")
                        dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
                        
                        with dl_col1:
                            txt_content = ""
                            if full_result.get("segments"):
                                txt_content += f"Transcription: {trans.get('filename', 'Unknown')}\n"
                                txt_content += f"Language: {trans.get('language', 'Unknown')}\n"
                                txt_content += f"Duration: {trans.get('duration', 0):.2f} seconds\n"
                                txt_content += f"Date: {trans.get('date', 'Unknown')}\n"
                                txt_content += "=" * 50 + "\n\n"
                                txt_content += "TRANSCRIPTION WITH TIMESTAMPS:\n"
                                txt_content += "-" * 50 + "\n\n"
                                for segment in full_result["segments"]:
                                    start = segment.get('start', 0)
                                    end = segment.get('end', 0)
                                    text = segment.get('text', '').strip()
                                    txt_content += f"[{format_timestamp(start)} --> {format_timestamp(end)}] {text}\n\n"
                                txt_content += "\n" + "=" * 50 + "\n\n"
                                txt_content += "PLAIN TEXT (NO TIMESTAMPS):\n"
                                txt_content += "-" * 50 + "\n\n"
                            txt_content += trans.get('text', '')
                            
                            st.download_button(
                                label="📄 TXT",
                                data=txt_content,
                                file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.txt",
                                mime="text/plain",
                                key=f"archived_txt_{trans.get('id')}"
                            )
                        
                        with dl_col2:
                            docx_buffer = create_docx(
                                trans.get('text', ''),
                                trans.get('filename', 'transcription'),
                                {
                                    "filename": trans.get('filename', 'Unknown'),
                                    "language": trans.get('language', 'Unknown'),
                                    "duration": trans.get('duration', 0),
                                    "date": trans.get('date', 'Unknown')
                                },
                                segments=full_result.get("segments")
                            )
                            st.download_button(
                                label="📝 DOCX",
                                data=docx_buffer.getvalue(),
                                file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"archived_docx_{trans.get('id')}"
                            )
                        
                        with dl_col3:
                            if trans.get('full_result'):
                                json_output = json.dumps(trans['full_result'], indent=2, ensure_ascii=False)
                                st.download_button(
                                    label="📊 JSON",
                                    data=json_output,
                                    file_name=f"{Path(trans.get('filename', 'transcription')).stem}_transcription.json",
                                    mime="application/json",
                                    key=f"archived_json_{trans.get('id')}"
                                )
                        
                        with dl_col4:
                            # Purge button for archived
                            if st.button("🗑️ Purge", key=f"purge_archived_{trans.get('id')}", help="Permanently delete", type="secondary"):
                                purge_transcription(trans.get('id'), from_archive=True)
                                st.success("✅ Transcription purged!")
                                st.rerun()
                        
                        if idx < len(filtered_archived) - 1:
                            st.divider()

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: gray;'>
        <p>Powered by <a href='https://github.com/openai/whisper' target='_blank'>OpenAI Whisper</a></p>
    </div>
    """,
    unsafe_allow_html=True
)

